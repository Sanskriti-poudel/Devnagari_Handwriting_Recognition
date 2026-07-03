"""
Export helpers for the Devanagari OCR web app.

Turns recognized Nepali text + the original scanned page(s) into downloadable
files:

  * build_docx(text)            -> editable .docx (opens in Word / Google Docs)
  * build_searchable_pdf(pages) -> a PDF that *looks* like the original scan but
                                   carries an invisible Unicode text layer, so the
                                   page becomes selectable / copyable / Ctrl+F-able.

Everything runs offline. python-docx is the only new dependency; the searchable
PDF is built with PyMuPDF (fitz), which the web app already uses to read PDFs.
"""
import io
import os

import cv2

# Reuse the project's cross-platform Devanagari font finder (Windows Nirmala/
# Mangal, Linux Lohit/Noto-Devanagari, ...). Needed so the invisible PDF text
# layer can encode Devanagari code points for search/copy.
try:
    from data.generate_synth import default_font_paths
except Exception:  # pragma: no cover - generate_synth import should not fail
    def default_font_paths():
        return []


def _devanagari_font_file():
    """Path to a Devanagari-capable TTF/OTF, or None if none is installed."""
    for p in default_font_paths():
        if os.path.exists(p):
            return p
    return None


# ----------------------------------------------------------------------------- docx
def build_docx(text):
    """Recognized text -> .docx bytes, one paragraph per line.

    Each run is tagged with a Devanagari font for the complex-script slot (w:cs)
    so Word/Google Docs render the Nepali glyphs correctly while staying fully
    editable.
    """
    from docx import Document
    from docx.shared import Pt
    from docx.oxml.ns import qn

    doc = Document()
    for raw_line in (text or "").split("\n"):
        para = doc.add_paragraph()
        run = para.add_run(raw_line)
        run.font.size = Pt(14)
        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.get_or_add_rFonts()
        # ascii/hAnsi cover Latin; cs (complex script) is what Word uses for Devanagari
        rfonts.set(qn("w:ascii"), "Noto Sans Devanagari")
        rfonts.set(qn("w:hAnsi"), "Noto Sans Devanagari")
        rfonts.set(qn("w:cs"), "Mangal")

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()


# ----------------------------------------------------------------------------- searchable pdf
def build_searchable_pdf(pages):
    """Build a searchable PDF from OCR'd pages.

    `pages` is a list of dicts:
        {"bgr": np.ndarray (H,W,3),
         "lines": [{"box": (x, y, w, h), "text": str}, ...]}

    For each page the original image is drawn at full size, then each line's
    recognized text is written on top in render mode 3 (invisible). The scan
    still looks identical, but the text is selectable and searchable.
    """
    import fitz  # PyMuPDF

    fontfile = _devanagari_font_file()
    doc = fitz.open()
    try:
        for pg in pages:
            bgr = pg["bgr"]
            h, w = bgr.shape[:2]
            page = doc.new_page(width=w, height=h)

            ok, buf = cv2.imencode(".png", bgr)
            if ok:
                page.insert_image(fitz.Rect(0, 0, w, h), stream=buf.tobytes())

            fontname = "deva" if fontfile else "helv"
            for ln in pg.get("lines", []):
                txt = (ln.get("text") or "").strip()
                if not txt:
                    continue
                x, y, bw, bh = ln["box"]
                rect = fitz.Rect(x, y, x + bw, y + bh)
                # size the (invisible) text to the line height, leaving room for
                # the font's ascent/descent so it fits inside a box this tall
                fontsize = max(6.0, min(float(bh) * 0.6, 48.0))
                placed = False
                try:
                    # insert_textbox returns the leftover vertical space; a NEGATIVE
                    # value means the text overflowed and NOTHING was drawn — that is
                    # not an exception, so we must check the return value explicitly.
                    rc = page.insert_textbox(
                        rect, txt,
                        fontsize=fontsize, fontname=fontname, fontfile=fontfile,
                        render_mode=3,   # invisible: visible scan stays, text is searchable
                    )
                    placed = rc >= 0
                except Exception:
                    placed = False
                if not placed:
                    # fall back to baseline placement; insert_text never clips, so the
                    # searchable layer is always written even for short/odd boxes
                    try:
                        page.insert_text(
                            (x, y + fontsize), txt,
                            fontsize=fontsize, fontname=fontname, fontfile=fontfile,
                            render_mode=3,
                        )
                    except Exception:
                        pass
        return doc.tobytes()
    finally:
        doc.close()
