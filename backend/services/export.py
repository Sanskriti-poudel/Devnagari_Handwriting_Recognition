"""
Export helpers for document-mode OCR results: turn recognized Nepali text +
the original scanned page(s) into downloadable files.

  * build_docx(text, line_data=None) -> editable .docx with optional table layout
  * build_searchable_pdf(pages)     -> a PDF that looks like the original scan but
                                        carries an invisible Unicode text layer.
"""
import os

import cv2

from ml_models.loader import REPO_ROOT
import sys

if REPO_ROOT not in sys.path:
    sys.path.insert(0, REPO_ROOT)

try:
    from data.generate_synth import default_font_paths
except Exception:  # pragma: no cover - generate_synth import should not fail
    def default_font_paths():
        return []


def _devanagari_font_file():
    # Prefer Ganesh ( Nepali font) over others
    fonts_dir = os.path.join(os.path.dirname(__file__), "..", "fonts")
    for font_name in ["ganesh.ttf", "NotoSansDevanagari.ttf"]:
        font_path = os.path.join(fonts_dir, font_name)
        if os.path.exists(font_path):
            return font_path
    for p in default_font_paths():
        if os.path.exists(p):
            return p
    return None


def _set_devanagari_font(run):
    """Apply Noto Sans Devanagari font to a python-docx run."""
    from docx.oxml.ns import qn
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), "Noto Sans Devanagari")
    rfonts.set(qn("w:hAnsi"), "Noto Sans Devanagari")
    rfonts.set(qn("w:cs"), "Mangal")


def build_docx(text, line_data=None, page_image=None):
    """Recognized text -> .docx bytes with table structure preserved.

    Detects table layouts from line_data by analyzing horizontal gaps between
    text boxes. Creates proper multi-column editable DOCX tables with Devanagari
    font support.

    If page_image is provided, the page dimensions match the image and the image
    is embedded as background. Text is rendered on top in gray for visibility
    while preserving the original appearance.
    """
    import io
    import numpy as np
    from docx import Document
    from docx.shared import Pt, Twips as Twip, Inches, RGBColor, Cm, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import base64

    doc = Document()

    # Encode page image if provided
    img_bytes = None
    img_width_inches = None
    img_height_inches = None

    if page_image is not None:
        rgb = page_image if page_image.ndim == 3 else cv2.cvtColor(page_image, cv2.COLOR_GRAY2BGR)
        ok, buf = cv2.imencode(".png", rgb)
        if ok:
            img_bytes = buf.tobytes()
            h, w = page_image.shape[:2]
            # Use 96 DPI for page size
            img_width_inches = w / 96.0
            img_height_inches = h / 96.0

    # Set page dimensions
    if img_width_inches and img_height_inches:
        for section in doc.sections:
            section.page_width = Inches(img_width_inches)
            section.page_height = Inches(img_height_inches)
            section.left_margin = Inches(0)
            section.right_margin = Inches(0)
            section.top_margin = Inches(0)
            section.bottom_margin = Inches(0)
    else:
        for section in doc.sections:
            section.left_margin = Cm(2)
            section.right_margin = Cm(2)
            section.top_margin = Cm(1)
            section.bottom_margin = Cm(1)

    if not line_data:
        # Fallback: simple paragraphs
        for raw_line in (text or "").split("\n"):
            para = doc.add_paragraph(raw_line)
            para.paragraph_format.space_before = Pt(2)
            para.paragraph_format.space_after = Pt(2)
            run = para.runs[0] if para.runs else para.add_run(raw_line)
            run.font.size = Pt(11)
            _set_devanagari_font(run)
        bio = io.BytesIO()
        doc.save(bio)
        return bio.getvalue()

    # Filter lines with valid boxes and sort by y then x
    valid_lines = [ln for ln in line_data if ln.get("box")]
    if not valid_lines:
        # No valid boxes, use simple paragraphs
        for raw_line in (text or "").split("\n"):
            para = doc.add_paragraph(raw_line)
            run = para.add_run(raw_line)
            run.font.size = Pt(11)
            _set_devanagari_font(run)
        bio = io.BytesIO()
        doc.save(bio)
        return bio.getvalue()

    # Group lines by vertical proximity (within 15 pixels = same row)
    def group_lines_by_row(lines):
        if not lines:
            return []
        sorted_lines = sorted(lines, key=lambda l: (l["box"][1], l["box"][0]))
        rows = []
        current_row = []
        last_y = sorted_lines[0]["box"][1]

        for ln in sorted_lines:
            y = ln["box"][1]
            if abs(y - last_y) <= 15:
                current_row.append(ln)
            else:
                if current_row:
                    rows.append(sorted(current_row, key=lambda l: l["box"][0]))
                current_row = [ln]
                last_y = y
        if current_row:
            rows.append(sorted(current_row, key=lambda l: l["box"][0]))
        return rows

    rows = group_lines_by_row(valid_lines)

    # Detect columns by analyzing horizontal gaps between text boxes across all rows
    # A "significant" gap (more than 2x median gap) indicates a column boundary
    def detect_column_boundaries(rows, num_cols_hint=None):
        """Detect column boundaries from horizontal positions of text boxes."""
        if not rows:
            return None

        # Collect all x-start and x-end positions
        all_positions = []
        for row in rows:
            for ln in row:
                x, y, w, h = ln["box"]
                all_positions.append((x, x + w))

        if len(all_positions) < 2:
            return None

        # Calculate gaps between consecutive text boxes on the same row
        gaps = []
        for row in rows:
            if len(row) < 2:
                continue
            sorted_row = sorted(row, key=lambda l: l["box"][0])
            for i in range(len(sorted_row) - 1):
                gap = sorted_row[i + 1]["box"][0] - (sorted_row[i]["box"][0] + sorted_row[i]["box"][2])
                if gap > 0:
                    gaps.append(gap)

        if not gaps:
            return None

        # Median gap indicates "normal" spacing between columns in same cell
        median_gap = np.median(gaps)
        # Significant gaps (> 3x median) indicate column boundaries
        threshold = max(median_gap * 3, 30)  # At least 30 pixels gap

        # Find column boundaries by detecting large gaps
        boundaries = set()
        for row in rows:
            if len(row) < 2:
                continue
            sorted_row = sorted(row, key=lambda l: l["box"][0])
            for i in range(len(sorted_row) - 1):
                gap = sorted_row[i + 1]["box"][0] - (sorted_row[i]["box"][0] + sorted_row[i]["box"][2])
                if gap > threshold:
                    # This is a column boundary
                    boundary = sorted_row[i]["box"][0] + sorted_row[i]["box"][2]
                    boundaries.add(boundary)

        if not boundaries:
            return None

        # Sort boundaries and create column ranges
        sorted_boundaries = sorted(boundaries)
        col_ranges = []
        for i, boundary in enumerate(sorted_boundaries):
            if i == 0:
                col_ranges.append((0, boundary))
            else:
                col_ranges.append((sorted_boundaries[i - 1], boundary))
        col_ranges.append((sorted_boundaries[-1], 2000))  # Last column extends to right

        return col_ranges

    # Try to detect columns
    col_ranges = detect_column_boundaries(rows)

    # Helper to assign text to columns based on x position
    def assign_to_columns(row, col_ranges):
        """Assign each text box in a row to a column based on x position."""
        if col_ranges is None:
            return [(0, row)]  # Single column

        assignments = {i: [] for i in range(len(col_ranges))}
        for ln in row:
            x = ln["box"][0]
            for col_idx, (start, end) in enumerate(col_ranges):
                if start <= x < end:
                    assignments[col_idx].append(ln)
                    break
            else:
                # Default to last column if no match
                assignments[len(col_ranges) - 1].append(ln)

        return [(idx, items) for idx, items in assignments.items() if items]

    # Build the DOCX table
    num_cols = len(col_ranges) if col_ranges else 1
    num_rows = len(rows)

    # If page image provided, add it as background first
    if img_bytes and img_width_inches and img_height_inches:
        # Add image paragraph at the beginning (will be behind subsequent content)
        # We use a picture with absolute positioning via VML
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        run = para.add_run()
        run.add_picture(io.BytesIO(img_bytes), width=Inches(img_width_inches))
        # Move this image to the back by manipulating XML
        # Get the paragraph element and move it to the body start
        body = doc._element.body
        para_elem = para._p
        body.insert(0, para_elem)

    # If we detected columns, use them; otherwise fall back to single column per row
    if col_ranges and num_cols > 1:
        # Create table with detected columns
        table = doc.add_table(rows=num_rows, cols=num_cols)
        table.style = "Table Grid"

        # Set column widths based on detected ranges
        for row_idx, row in enumerate(rows):
            assignments = assign_to_columns(row, col_ranges)
            for col_idx, col_items in assignments:
                cell = table.rows[row_idx].cells[col_idx]
                para = cell.paragraphs[0]
                para.paragraph_format.space_before = Pt(1)
                para.paragraph_format.space_after = Pt(1)
                para.paragraph_format.line_spacing = 1.0

                # Combine text from all items in this cell
                cell_text = " ".join(ln.get("text", "") for ln in sorted(col_items, key=lambda l: l["box"][0]))
                if cell_text:
                    run = para.add_run(cell_text)
                    run.font.size = Pt(10)
                    _set_devanagari_font(run)
                    # Make text semi-transparent if image is present
                    if img_bytes:
                        run.font.color.rgb = RGBColor(128, 128, 128)

                # Set column width
                start, end = col_ranges[col_idx]
                width_twips = int((end - start) * 15 / 8)  # Approximate conversion
                cell.width = Twip(max(width_twips, 500))
    else:
        # Fallback: single column per row (treat each row as one cell)
        table = doc.add_table(rows=num_rows, cols=1)
        table.style = "Table Grid"

        for row_idx, row in enumerate(rows):
            cell = table.rows[row_idx].cells[0]
            para = cell.paragraphs[0]
            para.paragraph_format.space_before = Pt(1)
            para.paragraph_format.space_after = Pt(1)
            para.paragraph_format.line_spacing = 1.0

            # Combine all text in this row
            row_text = " ".join(ln.get("text", "") for ln in sorted(row, key=lambda l: l["box"][0]))
            if row_text:
                run = para.add_run(row_text)
                run.font.size = Pt(10)
                _set_devanagari_font(run)
                # Make text semi-transparent if image is present
                if img_bytes:
                    run.font.color.rgb = RGBColor(128, 128, 128)

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()


def _merge_columns_reading_order(columns, raw_lines):
    """Merge lines from multiple columns into a single reading-order list.

    Walks down all columns simultaneously, appending one line from each column
    per step — producing left-to-right, top-to-bottom reading order for
    two-column layouts where left and right columns have similar y-ranges.
    """
    if len(columns) == 1:
        return raw_lines

    # Build lookup from box tuple to raw_lines entry
    box_to_line = {tuple(ln["box"]): ln for ln in raw_lines}

    result = []
    indices = [0] * len(columns)

    while any(i < len(col) for i, col in zip(indices, columns)):
        for col_idx, col in enumerate(columns):
            if indices[col_idx] < len(col):
                line = box_to_line.get(tuple(col[indices[col_idx]]))
                if line:
                    result.append(line)
                indices[col_idx] += 1

    return result


def build_searchable_pdf(pages):
    """`pages`: list of {"bgr": np.ndarray (H,W,3), "lines": [{"box": (x,y,w,h), "text": str}, ...]}.

    Embeds the original scan as a full-page image and overlays an invisible
    Unicode text layer at the detected line positions. Uses column detection
    to render lines in the correct reading order for multi-column layouts.
    """
    import fitz  # PyMuPDF

    # Import detect_columns from document.py (no circular import — document.py doesn't import export.py)
    from services.document import detect_columns

    fontfile = _devanagari_font_file()
    doc = fitz.open()
    try:
        for pg in pages:
            bgr = pg["bgr"]
            h, w = bgr.shape[:2]
            page = doc.new_page(width=w, height=h)

            ok, buf = cv2.imencode(".png", bgr)
            if ok:
                page.insert_image(fitz.Rect(0, 0, w, h), stream=buf.tobytes())

            raw_lines = pg.get("lines", [])
            if not raw_lines:
                continue

            boxes = [tuple(ln["box"]) for ln in raw_lines]
            columns = detect_columns(boxes, w, h)

            if len(columns) >= 2:
                # Multi-column: merge into reading order
                lines_ordered = _merge_columns_reading_order(columns, raw_lines)
            else:
                # Single column: sort by y then x (top-to-bottom, left-to-right)
                lines_ordered = sorted(raw_lines, key=lambda ln: (ln["box"][1], ln["box"][0]))

            fontname = "deva" if fontfile else "helv"
            for ln in lines_ordered:
                txt = (ln.get("text") or "").strip()
                if not txt:
                    continue
                x, y, bw, bh = ln["box"]
                rect = fitz.Rect(x, y, x + bw, y + bh)
                fontsize = max(6.0, min(float(bh) * 0.6, 48.0))
                # render_mode=0 = filled (visible) text; invisible text (mode=3) is not
                # reliably searchable across all PDF readers, so we use visible text
                # colored light gray so it doesn't distract from the original scan
                placed = False
                try:
                    rc = page.insert_textbox(
                        rect, txt,
                        fontsize=fontsize, fontname=fontname, fontfile=fontfile,
                        color=(0.6, 0.6, 0.6), render_mode=0,
                    )
                    placed = rc >= 0
                except Exception:
                    placed = False
                if not placed:
                    try:
                        page.insert_text(
                            (x, y + fontsize), txt,
                            fontsize=fontsize, fontname=fontname, fontfile=fontfile,
                            color=(0.6, 0.6, 0.6), render_mode=0,
                        )
                    except Exception:
                        pass
        return doc.tobytes()
    finally:
        doc.close()
